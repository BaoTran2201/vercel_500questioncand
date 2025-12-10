from __future__ import annotations
import json
import re
from pathlib import Path
from typing import List, Tuple, Optional
from uuid import uuid4

from docx import Document

DOC_PATH = Path(r"c:/Users/PC/Desktop/Road Traffic/500 câu hoi CAND.docx")
OUT_PATH = Path(r"c:/Users/PC/Desktop/Road Traffic/src/data/cand-questions.ts")
IMAGES_DIR = Path(r"c:/Users/PC/Desktop/Road Traffic/public/images")
IMAGES_DIR.mkdir(parents=True, exist_ok=True)

# Helper to pull text + underline flag from paragraph

def para_info(paragraph) -> Tuple[str, bool]:
    text = paragraph.text.strip()
    underlined = any(run.font.underline for run in paragraph.runs if run.text)
    return text, underlined


def is_question_start(text: str) -> bool:
    return bool(re.match(r"^Câu\s*\d+", text))


def extract_images_from_doc(doc: Document) -> dict:
    """Extract all images from document and save to IMAGES_DIR.
    Returns mapping of relationship ID to dict with url and size."""
    image_map = {}
    
    try:
        # Iterate through all relationships to find images
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                try:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # Get file extension from content type
                    content_type = image_part.content_type
                    ext_map = {
                        'image/jpeg': 'jpg',
                        'image/jpg': 'jpg',
                        'image/png': 'png',
                        'image/gif': 'gif',
                        'image/bmp': 'bmp',
                    }
                    ext = ext_map.get(content_type, 'png')
                    
                    # Generate filename
                    filename = f"question_{uuid4().hex[:8]}.{ext}"
                    image_path = IMAGES_DIR / filename
                    
                    # Save image
                    image_path.write_bytes(image_data)
                    image_map[rel_id] = {
                        "url": f"/images/{filename}",
                        "size": len(image_data),
                    }
                    print(f"  Extracted: {filename} ({content_type}, {len(image_data)} bytes)")
                except Exception as e:
                    print(f"  Warning: Failed to extract image: {e}")
    except Exception as e:
        print(f"  Warning: Failed to process document images: {e}")
    
    return image_map


def find_image_for_question(start_idx: int, end_idx: int, doc: Document, image_map: dict) -> Optional[str]:
    """Find image between start_idx (inclusive) and end_idx (exclusive).
    Collect all images and return the largest one (skip 648-byte placeholders)."""
    candidates = []
    
    try:
        for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
            para = doc.paragraphs[i]

            for run in para.runs:
                for drawing in run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                        embed = blip.get(embed_attr)
                        if not embed:
                            continue

                        # Look up from image_map which contains pre-extracted images
                        if embed in image_map:
                            img_info = image_map[embed]
                            # Only include images > 648 bytes (skip placeholders)
                            if img_info["size"] > 648:
                                candidates.append((img_info["size"], img_info["url"]))
    except Exception:
        pass

    # Return largest image found
    if candidates:
        candidates.sort(reverse=True)
        return candidates[0][1]
    return None


def is_question_start(text: str) -> bool:
    return bool(re.match(r"^Câu\s*\d+", text))


def parse_question_block(block: List[Tuple[str, bool]], question_id: int, start_para_idx: int, end_para_idx: int, doc: Document, image_map: dict):
    # block[0] is the question header line
    header_text = block[0][0]
    m = re.match(r"^Câu\s*(\d+)[:.]?\s*(.*)", header_text)
    if not m:
        return None
    qid = int(m.group(1))
    first_line_rest = m.group(2).strip()

    question_lines: List[str] = []
    if first_line_rest:
        question_lines.append(first_line_rest)

    # Split remaining lines into question continuation vs options
    idx = 1
    lines = block

    def option_starts(t: str) -> bool:
        return bool(
            re.match(r"^\d+\.\s", t)
            or re.match(r"^\d+$", t)
            or re.match(r"^\d+\.$", t)
            or t.startswith(".")
        )

    # collect question continuation
    while idx < len(lines):
        t, _ = lines[idx]
        if not t:
            idx += 1
            continue
        if option_starts(t) or is_question_start(t):
            break
        question_lines.append(t)
        idx += 1

    options: List[Tuple[str, bool]] = []
    current_text: List[str] = []
    current_under = False
    pending_queue: List[bool] = []  # underline info from number-only lines before the text appears

    def flush_current():
        nonlocal current_text, current_under
        if current_text:
            opt_text = " ".join(current_text).strip()
            options.append((opt_text, current_under))
            current_text = []
            current_under = False

    while idx < len(lines):
        t, under = lines[idx]
        idx += 1
        if not t:
            continue
        if is_question_start(t):
            # unexpected new question start inside block
            idx -= 1
            break
        num_match_inline = re.match(r"^(\d+)\.\s*(.*)$", t)
        num_only = re.match(r"^(\d+)\.?$", t)

        if num_match_inline:
            flush_current()
            current_text = [num_match_inline.group(2).strip()]
            use_under = pending_queue.pop(0) if pending_queue else under
            current_under = use_under or under
            continue
        if num_only:
            # Number-only line marks the start of the next option
            flush_current()
            pending_queue.append(bool(under))
            current_text = []
            current_under = False
            continue
        if t.startswith("."):
            flush_current()
            current_text = [t.lstrip(".").strip()]
            use_under = pending_queue.pop(0) if pending_queue else False
            current_under = use_under or under
            continue
        # continuation of current option text
        current_text.append(t)
        current_under = current_under or under

    flush_current()

    # Compute question string
    question = " ".join(question_lines).strip()

    answers = [opt for opt, _ in options]
    correct_idx = next((i for i, (_, u) in enumerate(options) if u), None)

    # Try to find an image for this question
    image_url = find_image_for_question(start_para_idx, end_para_idx, doc, image_map)

    result = {
        "id": qid,
        "question": question,
        "answers": answers,
        "correctAnswer": correct_idx,
    }
    
    if image_url:
        result["image"] = image_url
    
    return result


def parse_doc():
    doc = Document(DOC_PATH)
    
    print("Extracting images from document...")
    image_map = extract_images_from_doc(doc)
    
    blocks: List[List[Tuple[str, bool]]] = []
    para_indices: List[int] = []  # Track starting paragraph index for each block
    current: List[Tuple[str, bool]] = []
    start_idx = 0
    
    for para_idx, para in enumerate(doc.paragraphs):
        t, u = para_info(para)
        if not t:
            continue
        if is_question_start(t):
            if current:
                blocks.append(current)
                para_indices.append(start_idx)
            current = [(t, u)]
            start_idx = para_idx
        else:
            current.append((t, u))
    
    if current:
        blocks.append(current)
        para_indices.append(start_idx)

    parsed = []
    warnings = []
    for blk_idx, blk in enumerate(blocks):
        start_idx = para_indices[blk_idx]
        end_idx = para_indices[blk_idx + 1] if blk_idx + 1 < len(para_indices) else len(doc.paragraphs)
        q = parse_question_block(blk, blk_idx, start_idx, end_idx, doc, image_map)
        if not q:
            continue
        if q["correctAnswer"] is None:
            warnings.append(f"Question {q['id']} missing correct answer underline")
        parsed.append(q)

    return parsed, warnings


def to_typescript(questions):
    # Escape quotes in text, including smart quotes
    def esc(s: str) -> str:
        # Replace smart quotes with regular quotes using Unicode code points
        s = s.replace('\u201c', '"')  # Left double quotation mark
        s = s.replace('\u201d', '"')  # Right double quotation mark
        s = s.replace('\u2018', "'")  # Left single quotation mark
        s = s.replace('\u2019', "'")  # Right single quotation mark
        # Remove newlines and collapse whitespace
        s = ' '.join(s.split())
        # Then escape backslashes and regular quotes
        return s.replace('\\', '\\\\').replace('"', '\\"')

    lines = ["import type { Question } from './question-types';", "", "export const candQuestions: Question[] = ["]
    for q in questions:
        lines.append("  {")
        lines.append(f"    id: {q['id']},")
        lines.append(f"    question: \"{esc(q['question'])}\",")
        lines.append("    answers: [")
        for a in q['answers']:
            lines.append(f"      \"{esc(a)}\",")
        lines.append("    ],")
        correct = q['correctAnswer'] if q['correctAnswer'] is not None else -1
        lines.append(f"    correctAnswer: {correct},")
        if 'image' in q and q['image']:
            lines.append(f"    image: \"{q['image']}\",")
        lines.append("  },")
    lines.append("];\n")
    lines.append("export default candQuestions;\n")
    return "\n".join(lines)


def main():
    questions, warnings = parse_doc()
    OUT_PATH.write_text(to_typescript(questions), encoding="utf-8")
    print(f"\nParsed {len(questions)} questions -> {OUT_PATH}")
    images_count = sum(1 for q in questions if 'image' in q)
    print(f"Questions with images: {images_count}")
    if warnings:
        print("Warnings:")
        for w in warnings[:10]:
            print(" -", w)
        if len(warnings) > 10:
            print(f" ... and {len(warnings) - 10} more")


if __name__ == "__main__":
    main()
